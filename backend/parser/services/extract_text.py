from __future__ import annotations

import io
from typing import Literal

import pdfplumber
from docx import Document

AllowedFileTypes = Literal["pdf", "docx"]


def _get_ext(filename: str) -> str:
    return filename.split(".")[-1] if "." in filename else ""


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = _get_ext(filename).lower()
    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "docx":
        return _extract_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def _extract_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    text_parts = [paragraph.text for paragraph in document.paragraphs]

    # include cell text so tables do not get dropped entirely
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    return "\n".join(part for part in text_parts if part).strip()



